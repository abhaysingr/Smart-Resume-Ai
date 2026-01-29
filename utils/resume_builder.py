from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from io import BytesIO
import tempfile
import traceback

class ResumeBuilder:
    def __init__(self):
        self.templates = {
            "Modern": self.build_modern_template,
            "Professional": self.build_professional_template,
            "Minimal": self.build_minimal_template,
            "Creative": self.build_creative_template,
            "ATS-Friendly": self.build_ats_friendly_template
        }
        
    def collect_user_input(self):
        """
        Collect comprehensive user input for ATS-friendly resume
        Returns a structured dictionary with all resume data
        """
        data = {}
        
        print("\n" + "="*60)
        print("🎯 SMART RESUME AI - ATS-FRIENDLY RESUME BUILDER")
        print("="*60)
        
        # Personal Information (Main Header)
        print("\n📋 SECTION 1: PERSONAL INFORMATION")
        print("-" * 60)
        data['personal_info'] = {
            'full_name': input("Full Name: ").strip(),
            'email': input("Email ID (professional): ").strip(),
            'phone': input("Phone Number (with country code, e.g., +91-XXXXXXXXXX): ").strip(),
            'linkedin': input("LinkedIn URL (full URL): ").strip(),
            'github': input("GitHub URL (full URL): ").strip(),
            'location': input("Location (District, State, PIN Code, Country): ").strip()
        }
        
        # Professional Summary
        print("\n📝 SECTION 2: PROFESSIONAL SUMMARY")
        print("-" * 60)
        print("Enter 2-3 concise lines about your role, skills & career intent")
        data['summary'] = input("Professional Summary: ").strip()
        
        # Ask if AI should enhance the summary
        enhance = input("\n🤖 Would you like AI to enhance this summary? (yes/no): ").strip().lower()
        if enhance == 'yes':
            data['summary_needs_enhancement'] = True
        
        # Technical Skills
        print("\n💻 SECTION 3: TECHNICAL SKILLS")
        print("-" * 60)
        print("Enter skills separated by commas for each category")
        data['skills'] = {}
        
        categories = [
            ('programming_languages', 'Programming Languages'),
            ('frameworks_libraries', 'Frameworks & Libraries'),
            ('developer_tools', 'Developer Tools'),
            ('databases', 'Databases'),
            ('cloud_devops', 'Cloud & DevOps (if applicable)'),
            ('certifications', 'Additional Courses / Certifications (if any)')
        ]
        
        for key, label in categories:
            skills_input = input(f"{label}: ").strip()
            if skills_input:
                data['skills'][key] = [s.strip() for s in skills_input.split(',') if s.strip()]
        
        # Experience Type Selection
        print("\n💼 SECTION 4: EXPERIENCE")
        print("-" * 60)
        print("Choose your experience type:")
        print("1. Internship Experience (for freshers/students)")
        print("2. Professional Experience (for full-time work)")
        exp_type = input("Enter choice (1 or 2): ").strip()
        
        data['experience_type'] = 'internship' if exp_type == '1' else 'professional'
        data['experience'] = []
        
        num_exp = int(input(f"\nHow many {data['experience_type']} entries? "))
        
        for i in range(num_exp):
            print(f"\n--- Entry {i+1} ---")
            exp = {
                'position': input("Title/Position: ").strip(),
                'company': input("Company Name: ").strip(),
                'start_date': input("Start Date (Month Year, e.g., Jan 2023): ").strip(),
                'end_date': input("End Date (Month Year or 'Present'): ").strip(),
                'work_mode': input("Work Mode (Remote/Hybrid/On-site): ").strip(),
                'responsibilities': []
            }
            
            print("\nEnter 3 bullet points describing impact, tools & outcome:")
            for j in range(3):
                point = input(f"Point {j+1}: ").strip()
                exp['responsibilities'].append(point)
            
            # Ask if AI should enhance the points
            enhance = input("\n🤖 Would you like AI to enhance these points? (yes/no): ").strip().lower()
            if enhance == 'yes':
                exp['needs_enhancement'] = True
            
            data['experience'].append(exp)
        
        # Projects
        print("\n🛠️ SECTION 5: PROJECTS")
        print("-" * 60)
        data['projects'] = []
        
        num_projects = int(input("How many projects? "))
        
        for i in range(num_projects):
            print(f"\n--- Project {i+1} ---")
            project = {
                'name': input("Project Name: ").strip(),
                'technologies': input("Technologies Used (comma-separated): ").strip(),
                'github': input("GitHub URL (full link): ").strip(),
                'live_demo': input("Live Demo URL (optional, press Enter to skip): ").strip(),
                'responsibilities': []
            }
            
            print("\nEnter 3 bullet points (What/How/Why):")
            for j in range(3):
                point = input(f"Point {j+1}: ").strip()
                project['responsibilities'].append(point)
            
            # Ask if AI should enhance the points
            enhance = input("\n🤖 Would you like AI to enhance these points? (yes/no): ").strip().lower()
            if enhance == 'yes':
                project['needs_enhancement'] = True
            
            data['projects'].append(project)
        
        # Education
        print("\n🎓 SECTION 6: EDUCATION")
        print("-" * 60)
        data['education'] = []
        
        num_edu = int(input("How many education entries? "))
        
        for i in range(num_edu):
            print(f"\n--- Education {i+1} ---")
            edu = {
                'school': input("College/School Name: ").strip(),
                'location': input("Place, State: ").strip(),
                'degree': input("Degree (e.g., B.Tech in Computer Science): ").strip(),
                'field': input("Field of Study: ").strip(),
                'start_date': input("Start Date (Month Year): ").strip(),
                'graduation_date': input("End Date (Month Year): ").strip(),
                'gpa': input("GPA/Percentage (optional, press Enter to skip): ").strip()
            }
            data['education'].append(edu)
        
        # Template Selection
        print("\n🎨 SECTION 7: TEMPLATE SELECTION")
        print("-" * 60)
        print("Available Templates:")
        print("1. ATS-Friendly (Recommended for job applications)")
        print("2. Modern")
        print("3. Professional")
        print("4. Minimal")
        print("5. Creative")
        template_choice = input("Select template (1-5): ").strip()
        
        template_map = {
            '1': 'ATS-Friendly',
            '2': 'Modern',
            '3': 'Professional',
            '4': 'Minimal',
            '5': 'Creative'
        }
        data['template'] = template_map.get(template_choice, 'ATS-Friendly')
        
        print("\n✅ Data collection complete!")
        return data
    
    def enhance_with_ai(self, text, context="general"):
        """
        Placeholder for AI enhancement functionality
        In production, this would call an AI API (OpenAI, Claude, etc.)
        """
        # This is a placeholder - integrate your AI enhancement logic here
        print(f"🤖 AI Enhancement requested for: {context}")
        print(f"Original text: {text}")
        # Return original text for now - replace with actual AI call
        return text
    
    def generate_resume(self, data):
        """Generate a resume based on the provided data and template"""
        try:
            print(f"Starting resume generation with template: {data['template']}")
            
            # Create a new document
            doc = Document()
            
            # Select and apply template
            template_name = data['template'].lower()
            print(f"Using template: {template_name}")
            
            if template_name == 'ats-friendly':
                doc = self.build_ats_friendly_template(doc, data)
            elif template_name == 'modern':
                doc = self.build_modern_template(doc, data)
            elif template_name == 'professional':
                doc = self.build_professional_template(doc, data)
            elif template_name == 'minimal':
                doc = self.build_minimal_template(doc, data)
            elif template_name == 'creative':
                doc = self.build_creative_template(doc, data)
            else:
                print(f"Warning: Unknown template '{template_name}', falling back to ATS-friendly template")
                doc = self.build_ats_friendly_template(doc, data)
            
            # Save to buffer
            buffer = BytesIO()
            print("Saving document to buffer...")
            doc.save(buffer)
            buffer.seek(0)
            print("Resume generated successfully!")
            return buffer
            
        except Exception as e:
            print(f"Error in generate_resume: {str(e)}")
            print(f"Full traceback: {traceback.format_exc()}")
            print(f"Template data: {data}")
            raise

    def _format_list_items(self, items):
        """Helper function to handle both string and list inputs"""
        if isinstance(items, str):
            return [item.strip() for item in items.split('\n') if item.strip()]
        elif isinstance(items, list):
            return [item.strip() for item in items if item and item.strip()]
        return []

    def build_ats_friendly_template(self, doc, data):
        """
        Build ATS-Friendly resume template
        Optimized for Applicant Tracking Systems with clean, simple formatting
        """
        try:
            # Set up styles for ATS compatibility
            styles = doc.styles
            
            # Name style - Simple and bold
            name_style = styles.add_style('ATS Name', WD_STYLE_TYPE.PARAGRAPH) if 'ATS Name' not in styles else styles['ATS Name']
            name_style.font.size = Pt(16)
            name_style.font.bold = True
            name_style.font.color.rgb = RGBColor(0, 0, 0)
            name_style.font.name = 'Calibri'
            name_style.paragraph_format.space_after = Pt(2)
            name_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Contact style - Plain text
            contact_style = styles.add_style('ATS Contact', WD_STYLE_TYPE.PARAGRAPH) if 'ATS Contact' not in styles else styles['ATS Contact']
            contact_style.font.size = Pt(10)
            contact_style.font.name = 'Calibri'
            contact_style.font.color.rgb = RGBColor(0, 0, 0)
            contact_style.paragraph_format.space_after = Pt(2)
            contact_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Section header style - Bold and uppercase
            section_style = styles.add_style('ATS Section', WD_STYLE_TYPE.PARAGRAPH) if 'ATS Section' not in styles else styles['ATS Section']
            section_style.font.size = Pt(12)
            section_style.font.bold = True
            section_style.font.color.rgb = RGBColor(0, 0, 0)
            section_style.font.name = 'Calibri'
            section_style.paragraph_format.space_before = Pt(10)
            section_style.paragraph_format.space_after = Pt(4)

            # Normal text style
            normal_style = styles.add_style('ATS Normal', WD_STYLE_TYPE.PARAGRAPH) if 'ATS Normal' not in styles else styles['ATS Normal']
            normal_style.font.size = Pt(10)
            normal_style.font.name = 'Calibri'
            normal_style.font.color.rgb = RGBColor(0, 0, 0)
            normal_style.paragraph_format.space_after = Pt(2)

            # 1. MAIN HEADER - All in plain text
            name_para = doc.add_paragraph(data['personal_info']['full_name'].upper())
            name_para.style = name_style

            # Contact information - all in one line
            contact_parts = []
            if data['personal_info'].get('email'):
                contact_parts.append(data['personal_info']['email'])
            if data['personal_info'].get('phone'):
                contact_parts.append(data['personal_info']['phone'])
            if data['personal_info'].get('linkedin'):
                contact_parts.append(f"LinkedIn: {data['personal_info']['linkedin']}")
            if data['personal_info'].get('github'):
                contact_parts.append(f"GitHub: {data['personal_info']['github']}")
            if data['personal_info'].get('location'):
                contact_parts.append(data['personal_info']['location'])
            
            if contact_parts:
                contact = doc.add_paragraph(' | '.join(contact_parts))
                contact.style = contact_style

            # 2. PROFESSIONAL SUMMARY
            if data.get('summary'):
                doc.add_paragraph('PROFESSIONAL SUMMARY', style=section_style)
                summary = doc.add_paragraph(data['summary'])
                summary.style = normal_style
                summary.paragraph_format.space_after = Pt(8)

            # 3. TECHNICAL SKILLS
            if data.get('skills'):
                doc.add_paragraph('TECHNICAL SKILLS', style=section_style)
                skills = data['skills']
                
                skill_categories = [
                    ('programming_languages', 'Programming Languages'),
                    ('frameworks_libraries', 'Frameworks & Libraries'),
                    ('developer_tools', 'Developer Tools'),
                    ('databases', 'Databases'),
                    ('cloud_devops', 'Cloud & DevOps'),
                    ('certifications', 'Additional Courses / Certifications')
                ]
                
                for key, label in skill_categories:
                    if skills.get(key):
                        p = doc.add_paragraph()
                        p.style = normal_style
                        p.add_run(f"{label}: ").bold = True
                        p.add_run(', '.join(self._format_list_items(skills[key])))

            # 4. EXPERIENCE (Internship or Professional)
            if data.get('experience'):
                exp_header = 'INTERNSHIP EXPERIENCE' if data.get('experience_type') == 'internship' else 'PROFESSIONAL EXPERIENCE'
                doc.add_paragraph(exp_header, style=section_style)
                
                for exp in data['experience']:
                    # Position and Company
                    p = doc.add_paragraph()
                    p.style = normal_style
                    p.add_run(f"{exp['position']}").bold = True
                    p.add_run(f" | {exp['company']}")
                    
                    # Duration and Work Mode
                    duration = doc.add_paragraph()
                    duration.style = normal_style
                    duration.add_run(f"{exp['start_date']} – {exp['end_date']} | {exp.get('work_mode', 'On-site')}")
                    
                    # Bullet points
                    if exp.get('responsibilities'):
                        for resp in self._format_list_items(exp['responsibilities']):
                            bullet = doc.add_paragraph()
                            bullet.style = normal_style
                            bullet.paragraph_format.left_indent = Inches(0.25)
                            bullet.add_run('• ' + resp)
                    
                    # Add spacing after each experience entry
                    doc.add_paragraph().paragraph_format.space_after = Pt(6)

            # 5. PROJECTS
            if data.get('projects'):
                doc.add_paragraph('PROJECTS', style=section_style)
                
                for proj in data['projects']:
                    # Project name and technologies
                    p = doc.add_paragraph()
                    p.style = normal_style
                    p.add_run(proj['name']).bold = True
                    if proj.get('technologies'):
                        p.add_run(f" | {proj['technologies']}")
                    
                    # GitHub link
                    if proj.get('github'):
                        github_para = doc.add_paragraph(f"GitHub: {proj['github']}")
                        github_para.style = normal_style
                    
                    # Live demo (optional)
                    if proj.get('live_demo'):
                        demo_para = doc.add_paragraph(f"Live Demo: {proj['live_demo']}")
                        demo_para.style = normal_style
                    
                    # Bullet points
                    if proj.get('responsibilities'):
                        for resp in self._format_list_items(proj['responsibilities']):
                            bullet = doc.add_paragraph()
                            bullet.style = normal_style
                            bullet.paragraph_format.left_indent = Inches(0.25)
                            bullet.add_run('• ' + resp)
                    
                    # Add spacing after each project
                    doc.add_paragraph().paragraph_format.space_after = Pt(6)

            # 6. EDUCATION
            if data.get('education'):
                doc.add_paragraph('EDUCATION', style=section_style)
                
                for edu in data['education']:
                    # School name and location
                    p = doc.add_paragraph()
                    p.style = normal_style
                    school_text = edu['school']
                    if edu.get('location'):
                        school_text += f", {edu['location']}"
                    p.add_run(school_text).bold = True
                    
                    # Degree and field
                    degree_para = doc.add_paragraph()
                    degree_para.style = normal_style
                    degree_para.add_run(f"{edu['degree']} in {edu['field']}")
                    
                    # Duration
                    duration_para = doc.add_paragraph()
                    duration_para.style = normal_style
                    duration_text = f"Duration: {edu.get('start_date', '')} – {edu['graduation_date']}"
                    if edu.get('gpa'):
                        duration_text += f" | GPA: {edu['gpa']}"
                    duration_para.add_run(duration_text)
                    
                    # Add spacing after each education entry
                    doc.add_paragraph().paragraph_format.space_after = Pt(6)

            # Set standard margins for ATS compatibility
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.75)
                section.right_margin = Inches(0.75)

            return doc
            
        except Exception as e:
            print(f"Error in build_ats_friendly_template: {str(e)}")
            raise

    def build_modern_template(self, doc, data):
        """Build modern style resume with clean, minimalist design"""
        try:
            # Set up styles
            styles = doc.styles
            
            # Name style - Modern, clean look
            name_style = styles.add_style('Modern Name', WD_STYLE_TYPE.PARAGRAPH) if 'Modern Name' not in styles else styles['Modern Name']
            name_style.font.size = Pt(24)
            name_style.font.bold = True
            name_style.font.color.rgb = RGBColor(41, 128, 185)
            name_style.font.name = 'Arial'
            name_style.paragraph_format.space_after = Pt(0)
            name_style.paragraph_format.space_before = Pt(6)
            name_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Section style - Clean and modern
            section_style = styles.add_style('Modern Section', WD_STYLE_TYPE.PARAGRAPH) if 'Modern Section' not in styles else styles['Modern Section']
            section_style.font.size = Pt(14)
            section_style.font.bold = True
            section_style.font.color.rgb = RGBColor(41, 128, 185)
            section_style.font.name = 'Arial'
            section_style.paragraph_format.space_before = Pt(16)
            section_style.paragraph_format.space_after = Pt(4)

            # Section underline style
            section_underline = styles.add_style('Modern Section Underline', WD_STYLE_TYPE.PARAGRAPH) if 'Modern Section Underline' not in styles else styles['Modern Section Underline']
            section_underline.font.size = Pt(8)
            section_underline.font.color.rgb = RGBColor(41, 128, 185)
            section_underline.paragraph_format.space_after = Pt(8)

            # Normal text style
            normal_style = styles.add_style('Modern Normal', WD_STYLE_TYPE.PARAGRAPH) if 'Modern Normal' not in styles else styles['Modern Normal']
            normal_style.font.size = Pt(10)
            normal_style.font.name = 'Arial'
            normal_style.paragraph_format.space_after = Pt(2)
            normal_style.font.color.rgb = RGBColor(44, 62, 80)

            # Contact style
            contact_style = styles.add_style('Modern Contact', WD_STYLE_TYPE.PARAGRAPH) if 'Modern Contact' not in styles else styles['Modern Contact']
            contact_style.font.size = Pt(10)
            contact_style.font.name = 'Arial'
            contact_style.font.color.rgb = RGBColor(41, 128, 185)
            contact_style.paragraph_format.space_after = Pt(2)
            contact_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add name at the top
            name_paragraph = doc.add_paragraph(data['personal_info']['full_name'].upper())
            name_paragraph.style = name_style

            # Add role/title if available
            if data['personal_info'].get('title'):
                title = doc.add_paragraph(data['personal_info']['title'])
                title.style = contact_style

            # Contact information layout
            contact_info = doc.add_paragraph()
            contact_info.style = contact_style
            
            # Add contact details with separators
            contact_parts = []
            if data['personal_info'].get('email'): contact_parts.append(data['personal_info']['email'])
            if data['personal_info'].get('phone'): contact_parts.append(data['personal_info']['phone'])
            if data['personal_info'].get('location'): contact_parts.append(data['personal_info']['location'])
            if contact_parts:
                contact_info.add_run(' | '.join(contact_parts))

            # Links layout
            if data['personal_info'].get('linkedin') or data['personal_info'].get('portfolio') or data['personal_info'].get('github'):
                links = doc.add_paragraph()
                links.style = contact_style
                links_parts = []
                if data['personal_info'].get('linkedin'): links_parts.append(f"LinkedIn: {data['personal_info']['linkedin']}")
                if data['personal_info'].get('github'): links_parts.append(f"GitHub: {data['personal_info']['github']}")
                if data['personal_info'].get('portfolio'): links_parts.append(f"Portfolio: {data['personal_info']['portfolio']}")
                links.add_run(' | '.join(links_parts))

            # Professional Summary
            if data.get('summary'):
                doc.add_paragraph('PROFESSIONAL SUMMARY', style=section_style)
                doc.add_paragraph('_' * 40, style=section_underline)
                summary = doc.add_paragraph(data['summary'])
                summary.style = normal_style
                summary.paragraph_format.space_after = Pt(12)
                summary.paragraph_format.left_indent = Inches(0.2)

            # Experience Section
            if data.get('experience'):
                exp_header = 'INTERNSHIP EXPERIENCE' if data.get('experience_type') == 'internship' else 'EXPERIENCE'
                doc.add_paragraph(exp_header, style=section_style)
                doc.add_paragraph('_' * 40, style=section_underline)
                for exp in data['experience']:
                    p = doc.add_paragraph()
                    p.style = normal_style
                    p.paragraph_format.left_indent = Inches(0.2)
                    
                    # Company and position
                    p.add_run(f"{exp['position']} at {exp['company']}").bold = True
                    date_run = p.add_run(f"\n{exp['start_date']} - {exp['end_date']}")
                    date_run.font.color.rgb = RGBColor(41, 128, 185)
                    if exp.get('work_mode'):
                        work_mode_run = p.add_run(f" | {exp['work_mode']}")
                        work_mode_run.font.color.rgb = RGBColor(41, 128, 185)
                    
                    if exp.get('description'):
                        desc = doc.add_paragraph(exp['description'])
                        desc.style = normal_style
                        desc.paragraph_format.left_indent = Inches(0.4)
                    
                    if exp.get('responsibilities'):
                        for resp in self._format_list_items(exp['responsibilities']):
                            bullet = doc.add_paragraph()
                            bullet.style = normal_style
                            bullet.paragraph_format.left_indent = Inches(0.6)
                            bullet.add_run('• ' + resp)
                    p.paragraph_format.space_after = Pt(12)

            # Projects Section
            if data.get('projects'):
                doc.add_paragraph('PROJECTS', style=section_style)
                doc.add_paragraph('_' * 40, style=section_underline)
                for proj in data['projects']:
                    p = doc.add_paragraph()
                    p.style = normal_style
                    p.paragraph_format.left_indent = Inches(0.2)
                    
                    p.add_run(proj['name']).bold = True
                    if proj.get('technologies'):
                        tech_run = p.add_run(f" | {proj['technologies']}")
                        tech_run.font.color.rgb = RGBColor(41, 128, 185)
                    
                    if proj.get('github'):
                        github_para = doc.add_paragraph(f"GitHub: {proj['github']}")
                        github_para.style = normal_style
                        github_para.paragraph_format.left_indent = Inches(0.4)
                    
                    if proj.get('description'):
                        desc = doc.add_paragraph(proj['description'])
                        desc.style = normal_style
                        desc.paragraph_format.left_indent = Inches(0.4)
                    
                    if proj.get('responsibilities'):
                        for resp in self._format_list_items(proj['responsibilities']):
                            bullet = doc.add_paragraph()
                            bullet.style = normal_style
                            bullet.paragraph_format.left_indent = Inches(0.6)
                            bullet.add_run('• ' + resp)
                    p.paragraph_format.space_after = Pt(12)

            # Education Section
            if data.get('education'):
                doc.add_paragraph('EDUCATION', style=section_style)
                doc.add_paragraph('_' * 40, style=section_underline)
                for edu in data['education']:
                    p = doc.add_paragraph()
                    p.style = normal_style
                    p.paragraph_format.left_indent = Inches(0.2)
                    
                    school_text = edu['school']
                    if edu.get('location'):
                        school_text += f", {edu['location']}"
                    p.add_run(school_text).bold = True
                    p.add_run(f"\n{edu['degree']} in {edu['field']}")
                    date_run = p.add_run(f"\nDuration: {edu.get('start_date', '')} - {edu['graduation_date']}")
                    if edu.get('gpa'):
                        p.add_run(f" | GPA: {edu['gpa']}")
                    p.paragraph_format.space_after = Pt(8)

            # Skills Section
            if data.get('skills'):
                doc.add_paragraph('TECHNICAL SKILLS', style=section_style)
                doc.add_paragraph('_' * 40, style=section_underline)
                skills = data['skills']
                
                skill_categories = [
                    ('programming_languages', 'Programming Languages'),
                    ('frameworks_libraries', 'Frameworks & Libraries'),
                    ('developer_tools', 'Developer Tools'),
                    ('databases', 'Databases'),
                    ('cloud_devops', 'Cloud & DevOps'),
                    ('certifications', 'Certifications')
                ]
                
                for key, label in skill_categories:
                    if skills.get(key):
                        p = doc.add_paragraph()
                        p.style = normal_style
                        p.paragraph_format.left_indent = Inches(0.2)
                        p.add_run(f"{label}: ").bold = True
                        skills_text = ' • '.join(self._format_list_items(skills[key]))
                        p.add_run(skills_text)
                        p.paragraph_format.space_after = Pt(6)

            # Set margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.8)
                section.right_margin = Inches(0.8)

            return doc
            
        except Exception as e:
            print(f"Error in build_modern_template: {str(e)}")
            raise

    def build_professional_template(self, doc, data):
        """Build professional style resume with improved spacing and layout"""
        try:
            # Set up styles
            styles = doc.styles
            
            # Header style - Name
            header_style = styles.add_style('Pro Header', WD_STYLE_TYPE.PARAGRAPH) if 'Pro Header' not in styles else styles['Pro Header']
            header_style.font.size = Pt(24)
            header_style.font.bold = True
            header_style.font.color.rgb = RGBColor(0, 0, 0)
            header_style.paragraph_format.space_after = Pt(4)
            header_style.font.name = 'Calibri'

            # Section style
            section_style = styles.add_style('Pro Section', WD_STYLE_TYPE.PARAGRAPH) if 'Pro Section' not in styles else styles['Pro Section']
            section_style.font.size = Pt(14)
            section_style.font.bold = True
            section_style.font.color.rgb = RGBColor(0, 120, 215)
            section_style.paragraph_format.space_before = Pt(12)
            section_style.paragraph_format.space_after = Pt(6)
            section_style.font.name = 'Calibri'

            # Normal text style
            normal_style = styles.add_style('Pro Normal', WD_STYLE_TYPE.PARAGRAPH) if 'Pro Normal' not in styles else styles['Pro Normal']
            normal_style.font.size = Pt(10)
            normal_style.font.name = 'Calibri'
            normal_style.paragraph_format.space_after = Pt(2)

            # Contact style
            contact_style = styles.add_style('Pro Contact', WD_STYLE_TYPE.PARAGRAPH) if 'Pro Contact' not in styles else styles['Pro Contact']
            contact_style.font.size = Pt(10)
            contact_style.font.name = 'Calibri'
            contact_style.paragraph_format.space_after = Pt(6)

            # Add name at the top
            name_paragraph = doc.add_paragraph(data['personal_info']['full_name'])
            name_paragraph.style = header_style
            name_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Add contact information in a single line
            contact_parts = []
            if data['personal_info'].get('email'): contact_parts.append(data['personal_info']['email'])
            if data['personal_info'].get('phone'): contact_parts.append(data['personal_info']['phone'])
            if data['personal_info'].get('location'): contact_parts.append(data['personal_info']['location'])
            
            if contact_parts:
                contact = doc.add_paragraph()
                contact.style = contact_style
                contact.add_run(' | '.join(contact_parts))

            # Add LinkedIn and GitHub links
            links_parts = []
            if data['personal_info'].get('linkedin'): links_parts.append(f"LinkedIn: {data['personal_info']['linkedin']}")
            if data['personal_info'].get('github'): links_parts.append(f"GitHub: {data['personal_info']['github']}")
            if data['personal_info'].get('portfolio'): links_parts.append(f"Portfolio: {data['personal_info']['portfolio']}")
            
            if links_parts:
                links = doc.add_paragraph()
                links.style = contact_style
                links.add_run(' | '.join(links_parts))

            # Professional Summary
            if data.get('summary'):
                doc.add_paragraph('PROFESSIONAL SUMMARY', style=section_style)
                summary = doc.add_paragraph(data['summary'])
                summary.style = normal_style

            # Experience Section
            if data.get('experience'):
                exp_header = 'INTERNSHIP EXPERIENCE' if data.get('experience_type') == 'internship' else 'EXPERIENCE'
                doc.add_paragraph(exp_header, style=section_style)
                for exp in data['experience']:
                    p = doc.add_paragraph()
                    p.style = normal_style
                    p.add_run(f"{exp['position']} at {exp['company']}").bold = True
                    p.add_run(f" | {exp['start_date']} - {exp['end_date']}")
                    if exp.get('work_mode'):
                        p.add_run(f" | {exp['work_mode']}")
                    
                    if exp.get('description'):
                        desc = doc.add_paragraph(exp['description'])
                        desc.style = normal_style
                        desc.paragraph_format.left_indent = Inches(0.2)
                    
                    if exp.get('responsibilities'):
                        for resp in self._format_list_items(exp['responsibilities']):
                            bullet = doc.add_paragraph()
                            bullet.style = normal_style
                            bullet.paragraph_format.left_indent = Inches(0.3)
                            bullet.add_run('• ' + resp)

            # Projects Section
            if data.get('projects'):
                doc.add_paragraph('PROJECTS', style=section_style)
                for proj in data['projects']:
                    p = doc.add_paragraph()
                    p.style = normal_style
                    p.add_run(proj['name']).bold = True
                    if proj.get('technologies'):
                        p.add_run(f" | {proj['technologies']}")
                    if proj.get('github'):
                        p.add_run(f" | GitHub: {proj['github']}")
                    
                    if proj.get('description'):
                        desc = doc.add_paragraph(proj['description'])
                        desc.style = normal_style
                        desc.paragraph_format.left_indent = Inches(0.2)
                    
                    if proj.get('responsibilities'):
                        for resp in self._format_list_items(proj['responsibilities']):
                            bullet = doc.add_paragraph()
                            bullet.style = normal_style
                            bullet.paragraph_format.left_indent = Inches(0.3)
                            bullet.add_run('• ' + resp)

            # Education Section
            if data.get('education'):
                doc.add_paragraph('EDUCATION', style=section_style)
                for edu in data['education']:
                    p = doc.add_paragraph()
                    p.style = normal_style
                    school_text = edu['school']
                    if edu.get('location'):
                        school_text += f", {edu['location']}"
                    p.add_run(school_text).bold = True
                    p.add_run(f"\n{edu['degree']} in {edu['field']}")
                    p.add_run(f" | Duration: {edu.get('start_date', '')} - {edu['graduation_date']}")
                    if edu.get('gpa'):
                        p.add_run(f" | GPA: {edu['gpa']}")

            # Skills Section
            if data.get('skills'):
                doc.add_paragraph('TECHNICAL SKILLS', style=section_style)
                skills = data['skills']
                
                skill_categories = [
                    ('programming_languages', 'Programming Languages'),
                    ('frameworks_libraries', 'Frameworks & Libraries'),
                    ('developer_tools', 'Developer Tools'),
                    ('databases', 'Databases'),
                    ('cloud_devops', 'Cloud & DevOps'),
                    ('certifications', 'Certifications')
                ]
                
                for key, label in skill_categories:
                    if skills.get(key):
                        p = doc.add_paragraph()
                        p.style = normal_style
                        p.add_run(f"{label}: ").bold = True
                        skills_text = ', '.join(self._format_list_items(skills[key]))
                        p.add_run(skills_text)

            # Set margins for better space utilization
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.7)
                section.right_margin = Inches(0.7)

            return doc
            
        except Exception as e:
            print(f"Error in build_professional_template: {str(e)}")
            raise

    def build_minimal_template(self, doc, data):
        """Build minimal style resume"""
        try:
            # Set up styles
            styles = doc.styles
            
            # Header style - Large, bold name
            header_style = None
            if 'Min Header' not in styles:
                header_style = styles.add_style('Min Header', WD_STYLE_TYPE.PARAGRAPH)
                header_style.font.size = Pt(28)
                header_style.font.bold = True
                header_style.font.color.rgb = RGBColor(33, 33, 33)
                header_style.paragraph_format.space_after = Pt(4)
            else:
                header_style = styles['Min Header']
            
            # Contact style - Small, gray text
            contact_style = None
            if 'Min Contact' not in styles:
                contact_style = styles.add_style('Min Contact', WD_STYLE_TYPE.PARAGRAPH)
                contact_style.font.size = Pt(9)
                contact_style.font.color.rgb = RGBColor(100, 100, 100)
                contact_style.paragraph_format.space_after = Pt(12)
            else:
                contact_style = styles['Min Contact']
            
            # Section style - Medium, all caps
            section_style = None
            if 'Min Section' not in styles:
                section_style = styles.add_style('Min Section', WD_STYLE_TYPE.PARAGRAPH)
                section_style.font.size = Pt(12)
                section_style.font.all_caps = True
                section_style.font.bold = True
                section_style.font.color.rgb = RGBColor(33, 33, 33)
                section_style.paragraph_format.space_before = Pt(16)
                section_style.paragraph_format.space_after = Pt(8)
            else:
                section_style = styles['Min Section']
            
            # Normal text style
            normal_style = None
            if 'Min Normal' not in styles:
                normal_style = styles.add_style('Min Normal', WD_STYLE_TYPE.PARAGRAPH)
                normal_style.font.size = Pt(10)
                normal_style.font.color.rgb = RGBColor(33, 33, 33)
                normal_style.paragraph_format.space_after = Pt(4)
            else:
                normal_style = styles['Min Normal']
            
            # Add header with personal info
            personal = data['personal_info']
            name = doc.add_paragraph(personal['full_name'])
            name.style = header_style
            
            # Contact info in one line
            contact_parts = []
            if personal.get('email'): contact_parts.append(personal['email'])
            if personal.get('phone'): contact_parts.append(personal['phone'])
            if personal.get('location'): contact_parts.append(personal['location'])
            
            if contact_parts:
                contact = doc.add_paragraph()
                contact.style = contact_style
                contact.add_run(' • '.join(contact_parts))
            
            # Links in one line
            links_parts = []
            if personal.get('linkedin'): links_parts.append(f"LinkedIn: {personal['linkedin']}")
            if personal.get('github'): links_parts.append(f"GitHub: {personal['github']}")
            if personal.get('portfolio'): links_parts.append(f"Portfolio: {personal['portfolio']}")
            
            if links_parts:
                links = doc.add_paragraph()
                links.style = contact_style
                links.add_run(' • '.join(links_parts))
            
            # Professional Summary
            if data.get('summary'):
                doc.add_paragraph('SUMMARY', style=section_style)
                summary = doc.add_paragraph(data['summary'])
                summary.style = normal_style
            
            # Experience Section
            if data.get('experience'):
                exp_header = 'INTERNSHIP EXPERIENCE' if data.get('experience_type') == 'internship' else 'EXPERIENCE'
                doc.add_paragraph(exp_header, style=section_style)
                for exp in data['experience']:
                    p = doc.add_paragraph(style=normal_style)
                    p.add_run(f"{exp['position']} at {exp['company']}").bold = True
                    p.add_run(f"\n{exp['start_date']} - {exp['end_date']}")
                    if exp.get('work_mode'):
                        p.add_run(f" | {exp['work_mode']}")
                    
                    if exp.get('description'):
                        overview = doc.add_paragraph(exp['description'])
                        overview.style = normal_style
                    
                    if exp.get('responsibilities'):
                        for resp in self._format_list_items(exp['responsibilities']):
                            bullet = doc.add_paragraph(style=normal_style)
                            bullet.style.paragraph_format.left_indent = Inches(0.25)
                            bullet.add_run('• ' + resp)
            
            # Projects Section
            if data.get('projects'):
                doc.add_paragraph('PROJECTS', style=section_style)
                for proj in data['projects']:
                    p = doc.add_paragraph(style=normal_style)
                    p.add_run(proj['name']).bold = True
                    if proj.get('technologies'):
                        p.add_run(f"\nTechnologies: {proj['technologies']}")
                    if proj.get('github'):
                        p.add_run(f"\nGitHub: {proj['github']}")
                    
                    if proj.get('description'):
                        overview = doc.add_paragraph(proj['description'])
                        overview.style = normal_style
                    
                    if proj.get('responsibilities'):
                        for resp in self._format_list_items(proj['responsibilities']):
                            bullet = doc.add_paragraph(style=normal_style)
                            bullet.style.paragraph_format.left_indent = Inches(0.25)
                            bullet.add_run('• ' + resp)
            
            # Education Section
            if data.get('education'):
                doc.add_paragraph('EDUCATION', style=section_style)
                for edu in data['education']:
                    p = doc.add_paragraph(style=normal_style)
                    school_text = edu['school']
                    if edu.get('location'):
                        school_text += f", {edu['location']}"
                    p.add_run(f"{school_text} - {edu['degree']} in {edu['field']}").bold = True
                    p.add_run(f"\nDuration: {edu.get('start_date', '')} - {edu['graduation_date']}")
                    if edu.get('gpa'):
                        p.add_run(f" | GPA: {edu['gpa']}")
            
            # Skills Section
            if data.get('skills'):
                doc.add_paragraph('TECHNICAL SKILLS', style=section_style)
                skills = data['skills']
                
                skill_categories = [
                    ('programming_languages', 'Programming Languages'),
                    ('frameworks_libraries', 'Frameworks & Libraries'),
                    ('developer_tools', 'Developer Tools'),
                    ('databases', 'Databases'),
                    ('cloud_devops', 'Cloud & DevOps'),
                    ('certifications', 'Certifications')
                ]
                
                for key, label in skill_categories:
                    if skills.get(key):
                        p = doc.add_paragraph(style=normal_style)
                        p.add_run(f"{label}: ").bold = True
                        p.add_run(' • '.join(self._format_list_items(skills[key])))
            
            return doc
            
        except Exception as e:
            print(f"Error in build_minimal_template: {str(e)}")
            raise

    def build_creative_template(self, doc, data):
        """Build creative style resume with vibrant design and emojis"""
        try:
            # Set up styles
            styles = doc.styles
            
            # Name style - Creative and bold
            name_style = styles.add_style('Creative Name', WD_STYLE_TYPE.PARAGRAPH) if 'Creative Name' not in styles else styles['Creative Name']
            name_style.font.size = Pt(24)
            name_style.font.bold = True
            name_style.font.color.rgb = RGBColor(155, 89, 182)
            name_style.font.name = 'Arial'
            name_style.paragraph_format.space_after = Pt(4)
            name_style.paragraph_format.space_before = Pt(6)
            name_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Section style - Vibrant
            section_style = styles.add_style('Creative Section', WD_STYLE_TYPE.PARAGRAPH) if 'Creative Section' not in styles else styles['Creative Section']
            section_style.font.size = Pt(14)
            section_style.font.bold = True
            section_style.font.color.rgb = RGBColor(155, 89, 182)
            section_style.font.name = 'Arial'
            section_style.paragraph_format.space_before = Pt(16)
            section_style.paragraph_format.space_after = Pt(4)

            # Normal text style - Clean
            normal_style = styles.add_style('Creative Normal', WD_STYLE_TYPE.PARAGRAPH) if 'Creative Normal' not in styles else styles['Creative Normal']
            normal_style.font.size = Pt(10)
            normal_style.font.name = 'Arial'
            normal_style.paragraph_format.space_after = Pt(2)
            normal_style.font.color.rgb = RGBColor(52, 73, 94)

            # Contact style - Professional
            contact_style = styles.add_style('Creative Contact', WD_STYLE_TYPE.PARAGRAPH) if 'Creative Contact' not in styles else styles['Creative Contact']
            contact_style.font.size = Pt(10)
            contact_style.font.name = 'Arial'
            contact_style.font.color.rgb = RGBColor(155, 89, 182)
            contact_style.paragraph_format.space_after = Pt(2)
            contact_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add name at the top
            name_paragraph = doc.add_paragraph('✨ ' + data['personal_info']['full_name'] + ' ✨')
            name_paragraph.style = name_style

            # Add role/title if available
            if data['personal_info'].get('title'):
                title = doc.add_paragraph('💫 ' + data['personal_info']['title'])
                title.style = contact_style

            # Contact information layout
            contact_info = doc.add_paragraph()
            contact_info.style = contact_style
            
            contact_parts = []
            if data['personal_info'].get('email'): contact_parts.append(f"📧 {data['personal_info']['email']}")
            if data['personal_info'].get('phone'): contact_parts.append(f"📱 {data['personal_info']['phone']}")
            if data['personal_info'].get('location'): contact_parts.append(f"📍 {data['personal_info']['location']}")
            if contact_parts:
                contact_info.add_run(' | '.join(contact_parts))

            # Links with professional formatting
            if data['personal_info'].get('linkedin') or data['personal_info'].get('github') or data['personal_info'].get('portfolio'):
                links = doc.add_paragraph()
                links.style = contact_style
                links_parts = []
                if data['personal_info'].get('linkedin'): links_parts.append(f"🔗 LinkedIn: {data['personal_info']['linkedin']}")
                if data['personal_info'].get('github'): links_parts.append(f"🐙 GitHub: {data['personal_info']['github']}")
                if data['personal_info'].get('portfolio'): links_parts.append(f"🌐 Portfolio: {data['personal_info']['portfolio']}")
                links.add_run(' | '.join(links_parts))

            # Professional Summary
            if data.get('summary'):
                doc.add_paragraph('👨‍💼 PROFESSIONAL SUMMARY', style=section_style)
                summary = doc.add_paragraph(data['summary'])
                summary.style = normal_style
                summary.paragraph_format.space_after = Pt(12)
                summary.paragraph_format.left_indent = Inches(0.2)

            # Experience Section
            if data.get('experience'):
                exp_header = '🎓 INTERNSHIP EXPERIENCE' if data.get('experience_type') == 'internship' else '💼 EXPERIENCE'
                doc.add_paragraph(exp_header, style=section_style)
                for exp in data['experience']:
                    p = doc.add_paragraph()
                    p.style = normal_style
                    p.paragraph_format.left_indent = Inches(0.2)
                    
                    p.add_run(f"🚀 {exp['position']}").bold = True
                    p.add_run(f"\n🏢 {exp['company']}")
                    p.add_run(f"\n📅 {exp['start_date']} - {exp['end_date']}")
                    if exp.get('work_mode'):
                        p.add_run(f" | 📍 {exp['work_mode']}")
                    
                    if exp.get('description'):
                        desc = doc.add_paragraph(exp['description'])
                        desc.style = normal_style
                        desc.paragraph_format.left_indent = Inches(0.4)
                    
                    if exp.get('responsibilities'):
                        resp_para = doc.add_paragraph()
                        resp_para.style = normal_style
                        resp_para.paragraph_format.left_indent = Inches(0.4)
                        resp_para.add_run('🎯 Key Achievements:').bold = True
                        for resp in self._format_list_items(exp['responsibilities']):
                            bullet = doc.add_paragraph()
                            bullet.style = normal_style
                            bullet.paragraph_format.left_indent = Inches(0.6)
                            bullet.add_run('• ' + resp)
                    p.paragraph_format.space_after = Pt(12)

            # Projects Section
            if data.get('projects'):
                doc.add_paragraph('🛠️ PROJECTS', style=section_style)
                for proj in data['projects']:
                    p = doc.add_paragraph()
                    p.style = normal_style
                    p.paragraph_format.left_indent = Inches(0.2)
                    
                    p.add_run(f"✨ {proj['name']}").bold = True
                    if proj.get('technologies'):
                        p.add_run(f"\n💻 Technologies: {proj['technologies']}")
                    if proj.get('github'):
                        p.add_run(f"\n🐙 GitHub: {proj['github']}")
                    
                    if proj.get('description'):
                        desc = doc.add_paragraph(proj['description'])
                        desc.style = normal_style
                        desc.paragraph_format.left_indent = Inches(0.4)
                    
                    if proj.get('responsibilities'):
                        resp_para = doc.add_paragraph()
                        resp_para.style = normal_style
                        resp_para.paragraph_format.left_indent = Inches(0.4)
                        resp_para.add_run('🎯 Key Features:').bold = True
                        for resp in self._format_list_items(proj['responsibilities']):
                            bullet = doc.add_paragraph()
                            bullet.style = normal_style
                            bullet.paragraph_format.left_indent = Inches(0.6)
                            bullet.add_run('• ' + resp)
                    p.paragraph_format.space_after = Pt(12)

            # Education Section
            if data.get('education'):
                doc.add_paragraph('🎓 EDUCATION', style=section_style)
                for edu in data['education']:
                    p = doc.add_paragraph()
                    p.style = normal_style
                    p.paragraph_format.left_indent = Inches(0.2)
                    
                    school_text = edu['school']
                    if edu.get('location'):
                        school_text += f", {edu['location']}"
                    p.add_run(f"📚 {school_text}").bold = True
                    p.add_run(f"\n🎯 {edu['degree']} in {edu['field']}")
                    p.add_run(f"\n📅 Duration: {edu.get('start_date', '')} - {edu['graduation_date']}")
                    if edu.get('gpa'):
                        p.add_run(f" | 📊 GPA: {edu['gpa']}")
                    p.paragraph_format.space_after = Pt(8)

            # Skills Section
            if data.get('skills'):
                doc.add_paragraph('⭐ TECHNICAL SKILLS', style=section_style)
                skills = data['skills']
                
                skill_categories = [
                    ('programming_languages', 'Programming Languages', '💻'),
                    ('frameworks_libraries', 'Frameworks & Libraries', '📚'),
                    ('developer_tools', 'Developer Tools', '🛠️'),
                    ('databases', 'Databases', '🗄️'),
                    ('cloud_devops', 'Cloud & DevOps', '☁️'),
                    ('certifications', 'Certifications', '🏆')
                ]
                
                for key, label, icon in skill_categories:
                    if skills.get(key):
                        p = doc.add_paragraph()
                        p.style = normal_style
                        p.paragraph_format.left_indent = Inches(0.2)
                        p.add_run(f"{icon} {label}: ").bold = True
                        skills_text = ' • '.join(self._format_list_items(skills[key]))
                        p.add_run(skills_text)
                        p.paragraph_format.space_after = Pt(6)

            # Set margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.8)
                section.right_margin = Inches(0.8)

            return doc
            
        except Exception as e:
            print(f"Error in build_creative_template: {str(e)}")
            raise


def main():
    # Create builder instance
    builder = ResumeBuilder()
    
    # Collect user input
    print("Welcome to Smart Resume AI Builder!")
    print("This tool will help you create an ATS-friendly professional resume.\n")
    
    user_data = builder.collect_user_input()
    
    # Generate resume
    print("\n⏳ Generating your resume...")
    resume_buffer = builder.generate_resume(user_data)
    
    # Save to file
    output_filename = f"resume_{user_data['personal_info']['full_name'].replace(' ', '_')}.docx"
    with open(output_filename, 'wb') as f:
        f.write(resume_buffer.getvalue())
    
    print(f"\n✅ Resume successfully created: {output_filename}")
    print("Thank you for using Smart Resume AI Builder!")

if __name__ == "__main__":
    main()

